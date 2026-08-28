"""
Unit tests for document processing.
"""

import pytest
import io
import docx
import pypdf
from modules.document_processor import DocumentProcessor, DocumentProcessingError

def test_extract_from_text_success():
    sample = "This is a sample operational briefing."
    res = DocumentProcessor.extract_from_text(sample)
    assert res["source_type"] == "text"
    assert res["content"] == sample
    assert res["word_count"] == 6

def test_extract_from_text_empty_fails():
    with pytest.raises(DocumentProcessingError):
        DocumentProcessor.extract_from_text("   ")

def test_extract_from_txt_success():
    sample = b"Incident report content from text file."
    res = DocumentProcessor.extract_from_txt(sample, "test.txt")
    assert res["source_type"] == "txt"
    assert res["file_name"] == "test.txt"
    assert "Incident report" in res["content"]

def test_extract_from_docx_success():
    doc = docx.Document()
    doc.add_paragraph("First paragraph of cyber intelligence report.")
    doc.add_paragraph("Second paragraph with technical details.")
    
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    res = DocumentProcessor.extract_from_docx(file_bytes, "test.docx")
    assert res["source_type"] == "docx"
    assert "First paragraph" in res["content"]
    assert res["paragraph_count"] == 2

def test_process_upload_unsupported_format():
    with pytest.raises(DocumentProcessingError):
        DocumentProcessor.process_upload("document.exe", b"binary content")
