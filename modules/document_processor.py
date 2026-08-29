"""
Document Processor Module
Handles text extraction from pasted text, TXT, PDF, DOCX files,
images (OCR via pytesseract), and URLs (via trafilatura).
"""

import io
from typing import Dict, Any, Optional
import pypdf
import docx

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors."""
    pass

class DocumentProcessor:
    """Extracts clean text and metadata from supported document types."""

    @staticmethod
    def extract_from_text(text: str) -> Dict[str, Any]:
        """Process plain text input."""
        cleaned = text.strip()
        if not cleaned:
            raise DocumentProcessingError("Input text is empty.")
        
        return {
            "source_type": "text",
            "file_name": "Direct_Text_Input.txt",
            "char_count": len(cleaned),
            "word_count": len(cleaned.split()),
            "content": cleaned,
        }

    @staticmethod
    def extract_from_txt(file_bytes: bytes, file_name: str = "document.txt") -> Dict[str, Any]:
        """Extract text from TXT file bytes."""
        try:
            content = file_bytes.decode("utf-8", errors="replace").strip()
        except Exception as e:
            raise DocumentProcessingError(f"Failed to decode TXT file: {str(e)}")
            
        if not content:
            raise DocumentProcessingError("TXT file is empty.")
            
        return {
            "source_type": "txt",
            "file_name": file_name,
            "char_count": len(content),
            "word_count": len(content.split()),
            "content": content,
        }

    @staticmethod
    def extract_from_pdf(file_bytes: bytes, file_name: str = "document.pdf") -> Dict[str, Any]:
        """Extract text from PDF file bytes using PyPDF."""
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            num_pages = len(pdf_reader.pages)
            extracted_pages = []
            
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    extracted_pages.append(page_text.strip())
            
            content = "\n\n".join(extracted_pages).strip()
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse PDF document: {str(e)}")

        if not content:
            raise DocumentProcessingError("PDF does not contain any extractable text or is image-only.")

        return {
            "source_type": "pdf",
            "file_name": file_name,
            "page_count": num_pages,
            "char_count": len(content),
            "word_count": len(content.split()),
            "content": content,
        }

    @staticmethod
    def extract_from_docx(file_bytes: bytes, file_name: str = "document.docx") -> Dict[str, Any]:
        """Extract text from DOCX file bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
                        
            content = "\n\n".join(paragraphs).strip()
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse DOCX document: {str(e)}")

        if not content:
            raise DocumentProcessingError("DOCX document is empty.")

        return {
            "source_type": "docx",
            "file_name": file_name,
            "paragraph_count": len(paragraphs),
            "char_count": len(content),
            "word_count": len(content.split()),
            "content": content,
        }

    @staticmethod
    def extract_from_image(file_bytes: bytes, file_name: str = "image.png") -> Dict[str, Any]:
        """
        Extract text from image (JPG, PNG, WEBP) using OCR via pytesseract + Pillow.
        Requires Tesseract OCR to be installed system-wide.
        On Windows: https://github.com/UB-Mannheim/tesseract/wiki
        """
        try:
            from PIL import Image
        except ImportError:
            raise DocumentProcessingError(
                "Pillow is not installed. Run: pip install Pillow"
            )

        try:
            import pytesseract
        except ImportError:
            raise DocumentProcessingError(
                "pytesseract is not installed. Run: pip install pytesseract"
            )

        try:
            image = Image.open(io.BytesIO(file_bytes))
        except Exception as e:
            raise DocumentProcessingError(f"Could not open image file: {str(e)}")

        try:
            content = pytesseract.image_to_string(image, lang="eng").strip()
        except pytesseract.TesseractNotFoundError:
            raise DocumentProcessingError(
                "Tesseract OCR engine is not installed or not found on PATH. "
                "Please install it from: https://github.com/UB-Mannheim/tesseract/wiki — "
                "then restart the application."
            )
        except Exception as e:
            raise DocumentProcessingError(f"OCR extraction failed: {str(e)}")

        if not content:
            raise DocumentProcessingError(
                "No text could be extracted from the image. "
                "Ensure the image contains readable printed text and is not blurry or rotated."
            )

        return {
            "source_type": "image",
            "file_name": file_name,
            "char_count": len(content),
            "word_count": len(content.split()),
            "content": content,
        }

    @staticmethod
    def extract_from_url(url: str) -> Dict[str, Any]:
        """
        Scrape and extract clean article text from a public URL using trafilatura.
        Works with news articles, reports, blog posts, and most public pages.
        """
        url = url.strip()
        if not url:
            raise DocumentProcessingError("URL is empty.")

        if not (url.startswith("http://") or url.startswith("https://")):
            url = "https://" + url

        try:
            import trafilatura
        except ImportError:
            raise DocumentProcessingError(
                "trafilatura is not installed. Run: pip install trafilatura"
            )

        try:
            downloaded = trafilatura.fetch_url(url)
        except Exception as e:
            raise DocumentProcessingError(f"Failed to fetch URL: {str(e)}")

        if not downloaded:
            raise DocumentProcessingError(
                "Could not download page content. "
                "Check that the URL is public and accessible."
            )

        try:
            content = trafilatura.extract(
                downloaded,
                include_comments=False,
                include_tables=True,
                no_fallback=False,
            )
        except Exception as e:
            raise DocumentProcessingError(f"Content extraction from URL failed: {str(e)}")

        if not content or len(content.strip()) < 50:
            raise DocumentProcessingError(
                "Extracted content is too short or empty. "
                "This page may require login, JavaScript rendering, or block scraping."
            )

        content = content.strip()
        # Derive a clean filename from the URL domain
        try:
            from urllib.parse import urlparse
            domain = urlparse(url).netloc.replace("www.", "")
            file_name = f"Web_Article_{domain}.txt"
        except Exception:
            file_name = "Web_Article.txt"

        return {
            "source_type": "url",
            "file_name": file_name,
            "source_url": url,
            "char_count": len(content),
            "word_count": len(content.split()),
            "content": content,
        }

    @classmethod
    def process_upload(cls, file_name: str, file_bytes: bytes) -> Dict[str, Any]:
        """Router method to extract text based on file extension."""
        ext = file_name.lower().rsplit(".", 1)[-1]
        
        if ext == "txt":
            return cls.extract_from_txt(file_bytes, file_name)
        elif ext == "pdf":
            return cls.extract_from_pdf(file_bytes, file_name)
        elif ext in ["docx", "doc"]:
            return cls.extract_from_docx(file_bytes, file_name)
        elif ext in ["jpg", "jpeg", "png", "webp"]:
            return cls.extract_from_image(file_bytes, file_name)
        else:
            raise DocumentProcessingError(
                f"Unsupported file format: .{ext}. "
                "Please provide TXT, PDF, DOCX, JPG, PNG, or WEBP."
            )
